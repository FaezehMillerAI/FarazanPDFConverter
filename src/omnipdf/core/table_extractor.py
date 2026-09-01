"""
Table Extractor & Builder: Extracts tabular data and geometry from PDF pages
and reconstructs high-fidelity, styled Word (DOCX) tables.
"""

from typing import List, Dict, Any, Optional, Tuple
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import fitz


class TableExtractor:
    """Extracts tables from PDF pages and renders them into Word documents."""

    @classmethod
    def extract_page_tables(cls, page: fitz.Page) -> List[Dict[str, Any]]:
        """Extract all tables on a page using PyMuPDF table finder."""
        extracted = []
        try:
            tab_finder = page.find_tables()
            if not tab_finder or not tab_finder.tables:
                return []

            for tab in tab_finder.tables:
                bbox = tab.bbox  # (x0, y0, x1, y1)
                df_data = tab.extract()  # List of rows, each row is a list of cell strings
                if not df_data or len(df_data) == 0:
                    continue

                num_rows = len(df_data)
                num_cols = len(df_data[0]) if num_rows > 0 else 0
                if num_rows == 0 or num_cols == 0:
                    continue

                # Calculate approximate column widths relative to table width
                col_widths = []
                tab_width = bbox[2] - bbox[0]
                if tab.col_count > 0:
                    col_w = tab_width / tab.col_count
                    col_widths = [col_w] * tab.col_count
                else:
                    col_widths = [tab_width / num_cols] * num_cols

                extracted.append({
                    "bbox": bbox,
                    "rows": df_data,
                    "row_count": num_rows,
                    "col_count": num_cols,
                    "col_widths": col_widths,
                    "header_present": num_rows > 1,
                })
        except Exception as e:
            # Fallback if table finder fails
            return []

        return extracted

    @classmethod
    def render_table(
        cls,
        doc: docx.Document,
        table_dict: Dict[str, Any],
        caption: Optional[str] = None,
        style: str = "booktabs",
    ):
        """
        Render extracted table data into docx with professional academic/business styling.
        Supports 'booktabs' (LaTeX standard with top, header-mid, and bottom rules) or 'grid'.
        """
        rows = table_dict.get("rows", [])
        if not rows:
            return

        num_rows = len(rows)
        num_cols = len(rows[0])
        if num_rows == 0 or num_cols == 0:
            return

        # Add caption before table if provided
        if caption:
            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap.paragraph_format.space_before = Pt(6)
            p_cap.paragraph_format.space_after = Pt(4)
            r_cap = p_cap.add_run(caption)
            r_cap.bold = True
            r_cap.font.name = "Calibri"
            r_cap.font.size = Pt(9.5)
            r_cap.font.color.rgb = RGBColor(0x37, 0x41, 0x51)

        # Create DOCX table
        docx_table = doc.add_table(rows=num_rows, cols=num_cols)
        docx_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        docx_table.autofit = True

        # Calculate column widths in Inches (assuming total width ~ 6.5 inches max)
        total_pdf_width = sum(table_dict.get("col_widths", [1] * num_cols)) or 1.0
        max_table_inch = 6.5
        proportions = [w / total_pdf_width for w in table_dict.get("col_widths", [1] * num_cols)]

        for col_idx, col in enumerate(docx_table.columns):
            w_inch = max(0.6, proportions[col_idx] * max_table_inch)
            for cell in col.cells:
                cell.width = Inches(w_inch)

        # Style each cell
        for row_idx, row_data in enumerate(rows):
            is_header = (row_idx == 0 and num_rows > 1)
            for col_idx, cell_value in enumerate(row_data):
                if col_idx >= num_cols:
                    continue
                cell = docx_table.cell(row_idx, col_idx)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                # Cell padding
                pad_xml = f"""
                <w:tcMar {nsdecls("w")}>
                    <w:top w:w="80" w:type="dxa"/>
                    <w:left w:w="100" w:type="dxa"/>
                    <w:bottom w:w="80" w:type="dxa"/>
                    <w:right w:w="100" w:type="dxa"/>
                </w:tcMar>
                """
                cell._tc.get_or_add_tcPr().append(parse_xml(pad_xml))

                # Header styling
                if is_header:
                    shd_xml = f'<w:shd {nsdecls("w")} w:fill="F3F4F6" w:val="clear"/>'
                    cell._tc.get_or_add_tcPr().append(parse_xml(shd_xml))

                # Borders
                if style == "booktabs":
                    # Top line for row 0, double/heavy line for header bottom, single bottom for last row
                    top_sz = "12" if row_idx == 0 else "0"
                    bottom_sz = "8" if is_header else ("12" if row_idx == num_rows - 1 else "0")
                    top_val = "single" if top_sz != "0" else "none"
                    bottom_val = "single" if bottom_sz != "0" else "none"
                    border_xml = f"""
                    <w:tcBorders {nsdecls("w")}>
                        <w:top w:val="{top_val}" w:sz="{top_sz}" w:space="0" w:color="333333"/>
                        <w:left w:val="none"/>
                        <w:bottom w:val="{bottom_val}" w:sz="{bottom_sz}" w:space="0" w:color="333333"/>
                        <w:right w:val="none"/>
                    </w:tcBorders>
                    """
                    cell._tc.get_or_add_tcPr().append(parse_xml(border_xml))
                else:
                    # Grid borders
                    border_xml = f"""
                    <w:tcBorders {nsdecls("w")}>
                        <w:top w:val="single" w:sz="4" w:space="0" w:color="D1D5DB"/>
                        <w:left w:val="single" w:sz="4" w:space="0" w:color="D1D5DB"/>
                        <w:bottom w:val="single" w:sz="4" w:space="0" w:color="D1D5DB"/>
                        <w:right w:val="single" w:sz="4" w:space="0" w:color="D1D5DB"/>
                    </w:tcBorders>
                    """
                    cell._tc.get_or_add_tcPr().append(parse_xml(border_xml))

                # Paragraph content
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
                p.paragraph_format.line_spacing = 1.05

                val_clean = str(cell_value or "").strip()
                run = p.add_run(val_clean)
                run.font.name = "Calibri"
                run.font.size = Pt(9.0) if not is_header else Pt(9.5)
                run.bold = is_header
                if is_header:
                    run.font.color.rgb = RGBColor(0x11, 0x18, 0x27)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
                    # Align numbers right, text left
                    if val_clean.replace(".", "", 1).replace("-", "", 1).isdigit():
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    else:
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Spacing after table
        after_p = doc.add_paragraph()
        after_p.paragraph_format.space_before = Pt(4)
        after_p.paragraph_format.space_after = Pt(6)
