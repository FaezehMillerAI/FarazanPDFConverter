"""
Academic & Computer Science Paper Converter: Specializes in IEEE, ACM, arXiv,
Springer, NeurIPS formats with 2-column flow, native OMML math equations,
algorithm/code boxes, and structured tables/figures.
"""

import io
import os
import re
from typing import Optional, Callable, Dict, Any, List, Tuple
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from PIL import Image
import fitz

from omnipdf.core.layout_detector import LayoutDetector, LayoutBlock, BlockType, TextSpan
from omnipdf.core.math_engine import MathEngine
from omnipdf.core.code_detector import CodeDetector
from omnipdf.core.table_extractor import TableExtractor


class AcademicConverter:
    """High-fidelity converter for Computer Science, IEEE/ACM, and academic papers."""

    def __init__(
        self,
        pdf_path: str,
        docx_path: str,
        convert_math: bool = True,
        highlight_code: bool = True,
        extract_images: bool = True,
        image_dpi: int = 300,
        on_progress: Optional[Callable[[int, int, str], None]] = None,
    ):
        self.pdf_path = pdf_path
        self.docx_path = docx_path
        self.convert_math = convert_math
        self.highlight_code = highlight_code
        self.extract_images = extract_images
        self.image_dpi = image_dpi
        self.on_progress = on_progress

    def convert(self, page_range: Optional[Tuple[int, int]] = None) -> Dict[str, Any]:
        """Convert the PDF to a high-quality academic Word document."""
        if not os.path.exists(self.pdf_path):
            raise FileNotFoundError(f"PDF file not found: {self.pdf_path}")

        pdf_doc = fitz.open(self.pdf_path)
        total_pdf_pages = len(pdf_doc)

        start_page = page_range[0] if page_range else 0
        end_page = min(page_range[1], total_pdf_pages) if page_range else total_pdf_pages
        pages_to_process = range(start_page, end_page)
        num_pages = len(pages_to_process)

        doc = docx.Document()
        self._setup_document_styles(doc)

        # Set default page margins (0.75 inch)
        for section in doc.sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        stats = {
            "pages_processed": 0,
            "equations_converted": 0,
            "code_blocks": 0,
            "tables_extracted": 0,
            "figures_embedded": 0,
        }

        for page_idx_in_loop, page_idx in enumerate(pages_to_process):
            if self.on_progress:
                self.on_progress(page_idx_in_loop + 1, num_pages, f"Processing page {page_idx + 1}...")

            page = pdf_doc[page_idx]

            # 1. Extract tables
            extracted_tables = TableExtractor.extract_page_tables(page)
            stats["tables_extracted"] += len(extracted_tables)

            # 2. Geometric layout analysis and reading order
            layout_detector = LayoutDetector(page, page_num=page_idx + 1)
            ordered_blocks = layout_detector.analyze(extracted_tables)

            # 3. Render blocks into DOCX
            self._render_blocks(doc, page, ordered_blocks, stats)

            # Add page break if not last page
            if page_idx_in_loop < num_pages - 1:
                doc.add_page_break()

            stats["pages_processed"] += 1

        pdf_doc.close()

        # Save output document
        os.makedirs(os.path.dirname(os.path.abspath(self.docx_path)), exist_ok=True)
        doc.save(self.docx_path)

        return stats

    def _setup_document_styles(self, doc: docx.Document):
        """Configure academic document font styles."""
        styles = doc.styles
        
        # Normal Body Style
        normal_style = styles["Normal"]
        normal_style.font.name = "Times New Roman"
        normal_style.font.size = Pt(10)
        normal_style.font.color.rgb = RGBColor(0x11, 0x18, 0x27)
        normal_style.paragraph_format.line_spacing = 1.15
        normal_style.paragraph_format.space_after = Pt(3)

        # Heading 1
        if "Heading 1" in styles:
            h1 = styles["Heading 1"]
            h1.font.name = "Times New Roman"
            h1.font.size = Pt(12)
            h1.font.bold = True
            h1.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            h1.paragraph_format.space_before = Pt(10)
            h1.paragraph_format.space_after = Pt(4)

        # Heading 2
        if "Heading 2" in styles:
            h2 = styles["Heading 2"]
            h2.font.name = "Times New Roman"
            h2.font.size = Pt(10.5)
            h2.font.bold = True
            h2.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
            h2.paragraph_format.space_before = Pt(8)
            h2.paragraph_format.space_after = Pt(3)

    def _render_blocks(
        self,
        doc: docx.Document,
        page: fitz.Page,
        blocks: List[LayoutBlock],
        stats: Dict[str, int],
    ):
        """Render parsed LayoutBlocks into DOCX elements."""
        for b in blocks:
            b_type = b.block_type

            # Header / Footer
            if b_type in (BlockType.HEADER, BlockType.FOOTER):
                # Header/footer text can be rendered as subtle small text
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(b.text.replace("\n", " "))
                run.font.name = "Times New Roman"
                run.font.size = Pt(8.5)
                run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
                continue

            # Title
            if b_type == BlockType.TITLE:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(8)
                run = p.add_run(b.text.replace("\n", " "))
                run.font.name = "Times New Roman"
                run.font.size = Pt(16)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
                continue

            # Authors / Affiliations
            if b_type == BlockType.AUTHORS:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(8)
                run = p.add_run(b.text.replace("\n", "  "))
                run.font.name = "Times New Roman"
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
                continue

            # Abstract
            if b_type == BlockType.ABSTRACT:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.4)
                p.paragraph_format.right_indent = Inches(0.4)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(8)
                p.paragraph_format.line_spacing = 1.15
                
                # Check for "Abstract—" prefix
                clean_text = b.text.strip()
                abs_match = re.match(r"^(Abstract|ABSTRACT|Index Terms|Keywords)[:\s—\-]+(.*)", clean_text, re.DOTALL)
                if abs_match:
                    lead_label = abs_match.group(1)
                    body_text = abs_match.group(2)
                    r_label = p.add_run(f"{lead_label}—")
                    r_label.bold = True
                    r_label.font.name = "Times New Roman"
                    r_label.font.size = Pt(9.5)

                    r_body = p.add_run(body_text.replace("\n", " "))
                    r_body.font.name = "Times New Roman"
                    r_body.font.size = Pt(9.5)
                    r_body.italic = True
                else:
                    r = p.add_run(clean_text.replace("\n", " "))
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(9.5)
                    r.italic = True
                continue

            # Headings
            if b_type == BlockType.HEADING_1:
                p = doc.add_paragraph(style="Heading 1")
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(3)
                r = p.add_run(b.text.replace("\n", " "))
                r.bold = True
                continue

            if b_type == BlockType.HEADING_2:
                p = doc.add_paragraph(style="Heading 2")
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(2)
                r = p.add_run(b.text.replace("\n", " "))
                r.bold = True
                continue

            if b_type == BlockType.HEADING_3:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(2)
                r = p.add_run(b.text.replace("\n", " "))
                r.font.name = "Times New Roman"
                r.font.size = Pt(10)
                r.italic = True
                r.bold = True
                continue

            # Display Equations
            if b_type == BlockType.DISPLAY_EQUATION:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)

                if self.convert_math:
                    success = MathEngine.insert_equation_into_paragraph(p, b.text, is_display=True)
                    if success:
                        stats["equations_converted"] += 1
                    else:
                        r = p.add_run(b.text)
                        r.font.name = "Cambria Math"
                        r.italic = True
                else:
                    r = p.add_run(b.text)
                    r.font.name = "Cambria Math"
                    r.italic = True
                continue

            # Code / Algorithm Blocks
            if b_type == BlockType.CODE_BLOCK:
                lines = [l.text for l in b.lines]
                is_algo = b.metadata.get("is_algorithm", False)
                title = lines[0] if is_algo else None
                code_body = lines[1:] if is_algo else lines
                CodeDetector.render_code_box(doc, code_body, title=title, is_algorithm=is_algo)
                stats["code_blocks"] += 1
                continue

            # Tables
            if b_type == BlockType.TABLE:
                tab_data = b.metadata.get("table_data", {})
                TableExtractor.render_table(doc, tab_data)
                continue

            # Figures / Images
            if b_type == BlockType.FIGURE:
                if self.extract_images:
                    self._embed_figure(doc, page, b.bbox)
                    stats["figures_embedded"] += 1
                continue

            # Captions (Figure / Table captions)
            if b_type == BlockType.CAPTION:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(6)
                r = p.add_run(b.text.replace("\n", " "))
                r.font.name = "Times New Roman"
                r.font.size = Pt(9.0)
                r.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
                r.bold = True
                continue

            # References
            if b_type == BlockType.REFERENCES:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.3)
                p.paragraph_format.first_line_indent = Inches(-0.3)
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.05
                r = p.add_run(b.text.replace("\n", " "))
                r.font.name = "Times New Roman"
                r.font.size = Pt(9.0)
                continue

            # Standard Paragraphs (with inline math and formatting runs)
            self._render_paragraph_with_spans(doc, b, stats)

    def _render_paragraph_with_spans(self, doc: docx.Document, block: LayoutBlock, stats: Dict[str, int]):
        """Render a body paragraph preserving bold, italic, font sizes, and inline math."""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15

        for line in block.lines:
            for span in line.spans:
                stext = span.text
                if not stext:
                    continue

                # Check if span is an inline math formula
                if self.convert_math and span.is_math and len(stext.strip()) > 1:
                    inserted = MathEngine.insert_equation_into_paragraph(p, stext, is_display=False)
                    if inserted:
                        stats["equations_converted"] += 1
                        continue

                # Regular formatted text run
                run = p.add_run(stext)
                run.font.name = "Times New Roman" if not span.is_mono else "Consolas"
                run.font.size = Pt(max(8.0, min(span.size, 14.0)))
                run.bold = span.is_bold
                run.italic = span.is_italic

                # RGB Color preservation
                if span.color != 0:
                    r_val = (span.color >> 16) & 255
                    g_val = (span.color >> 8) & 255
                    b_val = span.color & 255
                    run.font.color.rgb = RGBColor(r_val, g_val, b_val)
                else:
                    run.font.color.rgb = RGBColor(0x11, 0x18, 0x27)

            # Add space between lines if wrapped
            if not line.text.endswith(" ") and not line.text.endswith("-"):
                p.add_run(" ")

    def _embed_figure(self, doc: docx.Document, page: fitz.Page, bbox: Tuple[float, float, float, float]):
        """Clip and embed a figure or diagram from the PDF at high resolution."""
        try:
            # Clip page area of figure
            rect = fitz.Rect(bbox)
            if rect.is_empty or rect.width < 20 or rect.height < 20:
                return

            # Render at 300 DPI
            zoom = self.image_dpi / 72.0
            mat = fitz.Matrix(zoom, zoom)
            pix = page.get_pixmap(matrix=mat, clip=rect, alpha=False)

            img_bytes = pix.tobytes("png")
            img_stream = io.BytesIO(img_bytes)

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(2)

            # Determine picture width (max 6.0 inches)
            fig_w_inch = min(6.0, max(1.5, rect.width / 72.0))
            doc.add_picture(img_stream, width=Inches(fig_w_inch))
        except Exception:
            pass
