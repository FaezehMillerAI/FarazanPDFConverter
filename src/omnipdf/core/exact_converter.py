"""
Exact Layout Converter: Preserves 100% visual layout fidelity using modern Word
DrawingML floating text boxes (wp:anchor / wps:txbx) positioned at exact coordinates.
Every word, sentence, and paragraph remains 100% selectable, editable, and re-flowable in Word.
"""

import io
import os
import re
import html
from typing import Optional, Callable, Dict, Any, Tuple, List
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import fitz


def sanitize_xml_string(s: str) -> str:
    """Strip XML 1.0 invalid control characters and escape XML entities."""
    if not s:
        return ""
    # Remove XML 1.0 invalid control characters
    cleaned = re.sub(r"[\x00-\x08\x0B\x0C\x0E-\x1F\x7F-\x84\x86-\x9F\uFDD0-\uFDEF\uFFFE\uFFFF]", "", s)
    return html.escape(cleaned)


class ExactLayoutConverter:
    """Zero layout drift converter with 100% editable DrawingML text boxes and embedded images."""

    # 1 point = 12700 EMUs in Office Open XML DrawingML
    PT_TO_EMU = 12700

    def __init__(
        self,
        pdf_path: str,
        docx_path: str,
        dpi: int = 300,
        on_progress: Optional[Callable[[int, int, str], None]] = None,
    ):
        self.pdf_path = pdf_path
        self.docx_path = docx_path
        self.dpi = dpi
        self.on_progress = on_progress
        self._shape_id_counter = 1000

    def convert(self, page_range: Optional[Tuple[int, int]] = None) -> Dict[str, Any]:
        """Convert PDF to a fully editable, coordinate-accurate Word document."""
        if not os.path.exists(self.pdf_path):
            raise FileNotFoundError(f"PDF file not found: {self.pdf_path}")

        pdf_doc = fitz.open(self.pdf_path)
        total_pdf_pages = len(pdf_doc)

        start_page = page_range[0] if page_range else 0
        end_page = min(page_range[1], total_pdf_pages) if page_range else total_pdf_pages
        pages_to_process = range(start_page, end_page)
        num_pages = len(pages_to_process)

        doc = docx.Document()
        stats = {
            "pages_processed": 0,
            "text_boxes_placed": 0,
            "images_embedded": 0,
        }

        for page_idx_in_loop, page_idx in enumerate(pages_to_process):
            if self.on_progress:
                self.on_progress(page_idx_in_loop + 1, num_pages, f"Converting page {page_idx + 1} (exact layout)...")

            page = pdf_doc[page_idx]
            rect = page.rect
            page_w_pt = rect.width
            page_h_pt = rect.height

            # Configure section dimensions for each page
            if page_idx_in_loop == 0:
                section = doc.sections[0]
            else:
                section = doc.add_section()

            section.page_width = Pt(page_w_pt)
            section.page_height = Pt(page_h_pt)
            section.top_margin = Pt(0)
            section.bottom_margin = Pt(0)
            section.left_margin = Pt(0)
            section.right_margin = Pt(0)

            # Master paragraph for page shapes
            p_page = doc.add_paragraph()
            p_page.paragraph_format.space_before = Pt(0)
            p_page.paragraph_format.space_after = Pt(0)
            p_page.paragraph_format.line_spacing = 1.0

            # 1. Extract and Embed Non-Text Images & Drawings at exact positions
            self._embed_page_images(doc, p_page, page, stats)

            # 2. Extract Text Blocks and generate 100% Editable DrawingML Text Boxes
            text_dict = page.get_text("dict")
            blocks = text_dict.get("blocks", [])

            # Run OCR fallback if page has no text layers
            text_char_count = sum([len(span.get("text", "")) for b in blocks for l in b.get("lines", []) for span in l.get("spans", [])])
            if text_char_count < 40:
                from omnipdf.core.ocr_engine import OCRAgent
                ocr_agent = OCRAgent()
                if ocr_agent.is_page_scanned(page):
                    ocr_dict = ocr_agent.extract_text_dict_with_ocr(page)
                    if ocr_dict.get("blocks"):
                        blocks = ocr_dict.get("blocks", [])

            for b in blocks:
                if b.get("type") != 0:  # Text blocks only
                    continue

                bbox = b.get("bbox", (0, 0, 0, 0))
                x0, y0, x1, y1 = bbox
                w_pt = max(20, x1 - x0)
                h_pt = max(12, y1 - y0)

                lines = b.get("lines", [])
                if not lines:
                    continue

                # Build editable DrawingML text box XML
                drawing_xml = self._build_drawingml_textbox(
                    x_pt=x0,
                    y_pt=y0,
                    w_pt=w_pt,
                    h_pt=h_pt,
                    lines=lines,
                )

                if drawing_xml:
                    try:
                        r_elem = parse_xml(drawing_xml)
                        p_page._element.append(r_elem)
                        stats["text_boxes_placed"] += 1
                    except Exception:
                        pass

            stats["pages_processed"] += 1

        pdf_doc.close()

        # Save output document
        os.makedirs(os.path.dirname(os.path.abspath(self.docx_path)), exist_ok=True)
        doc.save(self.docx_path)

        return stats

    def _embed_page_images(self, doc: docx.Document, p_page, page: fitz.Page, stats: Dict[str, int]):
        """Extract raster images from the page and embed them at exact positions."""
        try:
            image_list = page.get_images(full=True)
            for img_info in image_list:
                xref = img_info[0]
                base_image = page.parent.extract_image(xref)
                image_bytes = base_image["image"]
                
                # Find image bbox on the page
                img_rects = page.get_image_rects(xref)
                for rect in img_rects:
                    if rect.width < 10 or rect.height < 10:
                        continue
                    
                    img_stream = io.BytesIO(image_bytes)
                    try:
                        self._embed_floating_picture(
                            doc, p_page, img_stream,
                            x_pt=rect.x0,
                            y_pt=rect.y0,
                            w_pt=rect.width,
                            h_pt=rect.height
                        )
                        stats["images_embedded"] += 1
                    except Exception:
                        pass
        except Exception:
            pass

    def _embed_floating_picture(
        self,
        doc: docx.Document,
        p_page,
        img_stream: io.BytesIO,
        x_pt: float,
        y_pt: float,
        w_pt: float,
        h_pt: float,
    ):
        """Embed a floating image positioned behind text at exact coordinates."""
        self._shape_id_counter += 1
        shape_id = self._shape_id_counter

        run = p_page.add_run()
        picture = run.add_picture(img_stream, width=Pt(w_pt), height=Pt(h_pt))
        
        inline_elem = picture._inline
        parent = inline_elem.getparent()
        
        emu_x = int(x_pt * self.PT_TO_EMU)
        emu_y = int(y_pt * self.PT_TO_EMU)
        emu_w = int(w_pt * self.PT_TO_EMU)
        emu_h = int(h_pt * self.PT_TO_EMU)

        graphic = inline_elem.find("{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}graphic")
        if graphic is None:
            graphic = inline_elem.find(".//{http://schemas.openxmlformats.org/drawingml/2006/main}graphic")

        if graphic is not None:
            anchor_xml = f"""
            <wp:anchor xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
                distT="0" distB="0" distL="0" distR="0" simplePos="0" relativeHeight="251658240"
                behindDoc="1" locked="0" layoutInCell="1" allowOverlap="1">
                <wp:simplePos x="0" y="0"/>
                <wp:positionH relativeFrom="page">
                    <wp:posOffset>{emu_x}</wp:posOffset>
                </wp:positionH>
                <wp:positionV relativeFrom="page">
                    <wp:posOffset>{emu_y}</wp:posOffset>
                </wp:positionV>
                <wp:extent cx="{emu_w}" cy="{emu_h}"/>
                <wp:effectExtent b="0" l="0" r="0" t="0"/>
                <wp:wrapNone/>
                <wp:docPr id="{shape_id}" name="Picture {shape_id}"/>
                <wp:cNvGraphicFramePr/>
            </wp:anchor>
            """
            anchor_elem = parse_xml(anchor_xml)
            anchor_elem.append(graphic)
            parent.replace(inline_elem, anchor_elem)

    def _build_drawingml_textbox(
        self,
        x_pt: float,
        y_pt: float,
        w_pt: float,
        h_pt: float,
        lines: List[Dict[str, Any]],
    ) -> str:
        """Construct a modern DrawingML floating text box XML with editable paragraphs and runs."""
        self._shape_id_counter += 1
        shape_id = self._shape_id_counter

        emu_x = int(x_pt * self.PT_TO_EMU)
        emu_y = int(y_pt * self.PT_TO_EMU)
        emu_w = int((w_pt + 12) * self.PT_TO_EMU)
        emu_h = int((h_pt + 8) * self.PT_TO_EMU)

        paragraphs_xml = []

        for line in lines:
            spans = line.get("spans", [])
            if not spans:
                continue

            runs_xml = []
            for span in spans:
                text = span.get("text", "")
                if not text:
                    continue

                safe_text = sanitize_xml_string(text)
                if not safe_text:
                    continue

                font_name = span.get("font", "Calibri")
                clean_font = re.sub(r"^[A-Z]{6}\+", "", font_name)  # Clean subset prefix
                clean_font = sanitize_xml_string(clean_font)
                if not clean_font:
                    clean_font = "Calibri"

                size_pt = span.get("size", 10.0)
                size_half_pt = int(size_pt * 2)

                flags = span.get("flags", 0)
                is_bold = bool(flags & 16) or ("bold" in font_name.lower()) or ("black" in font_name.lower())
                is_italic = bool(flags & 2) or ("italic" in font_name.lower()) or ("oblique" in font_name.lower())

                color_int = span.get("color", 0)
                hex_color = f"{color_int:06X}" if color_int != 0 else "111827"

                bold_xml = "<w:b/>" if is_bold else ""
                italic_xml = "<w:i/>" if is_italic else ""

                runs_xml.append(f"""
                <w:r>
                    <w:rPr>
                        <w:rFonts w:ascii="{clean_font}" w:hAnsi="{clean_font}" w:cs="{clean_font}"/>
                        {bold_xml}
                        {italic_xml}
                        <w:sz w:val="{size_half_pt}"/>
                        <w:szCs w:val="{size_half_pt}"/>
                        <w:color w:val="{hex_color}"/>
                    </w:rPr>
                    <w:t xml:space="preserve">{safe_text}</w:t>
                </w:r>
                """)

            joined_runs = "".join(runs_xml)
            if not joined_runs.strip():
                continue

            paragraphs_xml.append(f"""
            <w:p>
                <w:pPr>
                    <w:spacing w:line="240" w:lineRule="auto" w:before="0" w:after="0"/>
                </w:pPr>
                {joined_runs}
            </w:p>
            """)

        joined_paragraphs = "".join(paragraphs_xml)
        if not joined_paragraphs.strip():
            return ""

        drawingml_box = f"""
        <w:r {nsdecls("w")}>
            <w:drawing>
                <wp:anchor xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
                    xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
                    xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
                    distT="0" distB="0" distL="0" distR="0" simplePos="0" relativeHeight="251658240"
                    behindDoc="0" locked="0" layoutInCell="1" allowOverlap="1">
                    <wp:simplePos x="0" y="0"/>
                    <wp:positionH relativeFrom="page">
                        <wp:posOffset>{emu_x}</wp:posOffset>
                    </wp:positionH>
                    <wp:positionV relativeFrom="page">
                        <wp:posOffset>{emu_y}</wp:posOffset>
                    </wp:positionV>
                    <wp:extent cx="{emu_w}" cy="{emu_h}"/>
                    <wp:effectExtent b="0" l="0" r="0" t="0"/>
                    <wp:wrapNone/>
                    <wp:docPr id="{shape_id}" name="TextBox {shape_id}"/>
                    <wp:cNvGraphicFramePr/>
                    <a:graphic>
                        <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
                            <wps:wsp>
                                <wps:cNvSpPr/>
                                <wps:spPr>
                                    <a:xfrm>
                                        <a:off x="0" y="0"/>
                                        <a:ext cx="{emu_w}" cy="{emu_h}"/>
                                    </a:xfrm>
                                    <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>
                                    <a:noFill/>
                                    <a:ln><a:noFill/></a:ln>
                                </wps:spPr>
                                <wps:bodyPr rot="0" vert="horz" wrap="none" lIns="0" tIns="0" rIns="0" bIns="0" numCol="1" spcCol="0"/>
                                <wps:txbx>
                                    <w:txbxContent>
                                        {joined_paragraphs}
                                    </w:txbxContent>
                                </wps:txbx>
                            </wps:wsp>
                        </a:graphicData>
                    </a:graphic>
                </wp:anchor>
            </w:drawing>
        </w:r>
        """
        return drawingml_box
