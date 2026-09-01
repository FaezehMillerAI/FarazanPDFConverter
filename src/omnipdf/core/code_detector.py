"""
Code and Algorithm Detector: Identifies programming code, pseudocode, algorithms,
and formats them into shaded, monospace callout boxes in Word (DOCX).
"""

import re
from typing import List, Dict, Any, Optional, Tuple
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn


class CodeDetector:
    """Detects, extracts, and renders code and algorithm listings in Word documents."""

    MONOSPACE_FONTS = {
        "courier", "couriernew", "consolas", "monaco", "menlo",
        "dejavusansmono", "inconsolata", "sourcecodepro", "firacode",
        "firamono", "cmtt", "typewriter", "monospace", "lucidaconsole"
    }

    KEYWORD_PATTERNS = [
        re.compile(r"^\s*(def|class|import|from|return|async|await|lambda)\b"),
        re.compile(r"^\s*(public|private|protected|static|final|void|class|interface)\b"),
        re.compile(r"^\s*(int|float|double|char|bool|string|void|auto|let|const|var)\b"),
        re.compile(r"^\s*(for|while|if|else|elif|switch|case|break|continue|try|catch|finally)\b"),
        re.compile(r"^\s*(SELECT|INSERT|UPDATE|DELETE|FROM|WHERE|JOIN|GROUP BY)\b", re.IGNORECASE),
        re.compile(r"^\s*Algorithm\s+\d+[:\.]?", re.IGNORECASE),
        re.compile(r"^\s*(Input|Output|Ensure|Require|Procedure|Function|Returns)[:\s]", re.IGNORECASE),
        re.compile(r"^\s*\d+:\s+"),  # Line numbered code e.g. "1: while True do"
    ]

    @classmethod
    def is_monospace_font(cls, font_name: str) -> bool:
        """Check if font family name indicates a monospace code font."""
        if not font_name:
            return False
        clean_name = re.sub(r"[^a-zA-Z0-9]", "", font_name.lower())
        return any(mono in clean_name for mono in cls.MONOSPACE_FONTS)

    @classmethod
    def is_code_line(cls, text: str, font_name: str = "") -> bool:
        """Heuristic check if a line is part of a code snippet or algorithm."""
        if cls.is_monospace_font(font_name):
            return True

        # Check keyword syntax
        for pattern in cls.KEYWORD_PATTERNS:
            if pattern.search(text):
                return True

        # Check symbol / syntax density
        if len(text.strip()) > 4 and (
            text.strip().endswith(";")
            or text.strip().endswith("{")
            or text.strip().endswith("}")
            or text.strip().startswith("//")
            or text.strip().startswith("/*")
            or text.strip().startswith("#")
            or re.search(r"[a-zA-Z0-9_]+\([a-zA-Z0-9_,\s]*\);?", text)
        ):
            return True

        return False

    @classmethod
    def is_code_block(cls, lines: List[Dict[str, Any]]) -> Tuple[bool, bool]:
        """
        Check if a group of lines forms a code block or algorithm.
        Returns (is_code, is_algorithm).
        """
        if not lines:
            return False, False

        code_score = 0
        is_algo = False

        for line_data in lines:
            text = line_data.get("text", "")
            font = line_data.get("font", "")

            if re.search(r"^\s*Algorithm\s+\d+[:\.]?", text, re.IGNORECASE):
                is_algo = True
                code_score += 3
            elif cls.is_monospace_font(font):
                code_score += 2
            elif cls.is_code_line(text, font):
                code_score += 1

        is_code = (code_score / max(1, len(lines))) >= 0.7 or is_algo
        return is_code, is_algo

    @classmethod
    def render_code_box(
        cls,
        doc: docx.Document,
        code_lines: List[str],
        title: Optional[str] = None,
        is_algorithm: bool = False,
    ):
        """
        Render code or algorithm in a shaded callout box in python-docx with
        clean styling, light background, borders, and preserved indentation.
        """
        # Create a single-cell container table
        table = doc.add_table(rows=1, cols=1)
        table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        cell = table.cell(0, 0)
        
        # Set background shading: #F6F8FA (GitHub style) or #F8FAFC
        bg_color = "F6F8FA" if not is_algorithm else "FAFAFA"
        shading_xml = f'<w:shd {nsdecls("w")} w:fill="{bg_color}" w:val="clear"/>'
        cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

        # Set borders: Left accent border (#0969DA or #64748B) and subtle outer border
        border_color = "0969DA" if not is_algorithm else "475569"
        borders_xml = f"""
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="single" w:sz="4" w:space="0" w:color="D0D7DE"/>
            <w:left w:val="single" w:sz="24" w:space="0" w:color="{border_color}"/>
            <w:bottom w:val="single" w:sz="4" w:space="0" w:color="D0D7DE"/>
            <w:right w:val="single" w:sz="4" w:space="0" w:color="D0D7DE"/>
        </w:tcBorders>
        """
        cell._tc.get_or_add_tcPr().append(parse_xml(borders_xml))

        # Set cell padding / margins
        mar_xml = f"""
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="120" w:type="dxa"/>
            <w:left w:w="160" w:type="dxa"/>
            <w:bottom w:w="120" w:type="dxa"/>
            <w:right w:w="160" w:type="dxa"/>
        </w:tcMar>
        """
        cell._tc.get_or_add_tcPr().append(parse_xml(mar_xml))

        # Clear default paragraph
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.15

        # If title / algorithm header exists
        if title:
            p_title = cell.paragraphs[0]
            r_title = p_title.add_run(title)
            r_title.bold = True
            r_title.font.name = "Calibri"
            r_title.font.size = Pt(10)
            r_title.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
            p_title.paragraph_format.space_after = Pt(4)
            # Add horizontal divider below title
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)

        for idx, line in enumerate(code_lines):
            if idx > 0 or title:
                p = cell.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(1)
                p.paragraph_format.line_spacing = 1.1

            run = p.add_run(line)
            run.font.name = "Consolas"
            run.font.size = Pt(9.0)
            run.font.color.rgb = RGBColor(0x24, 0x29, 0x2F)

        # Add empty spacing after table
        after_p = doc.add_paragraph()
        after_p.paragraph_format.space_before = Pt(4)
        after_p.paragraph_format.space_after = Pt(4)
